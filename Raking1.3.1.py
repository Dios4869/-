import os
import openpyxl
from docx import Document
from docx.shared import Inches

# 定义模板目录路径
TEMPLATES_DIR = "模板/"

# 定义输出目录路径
OUTPUT_DIR = "输出/"

# 定义模板文件名和对应的输出文件名
TEMPLATE_FILENAMES = {
    "（启源）（业主）（监督）项目备案登记表.docx": "（启源）（业主）（监督）项目备案登记表.docx",
    "（启源）（业主）招标文件编制批准表.docx": "（启源）（业主）招标文件编制批准表.docx",
    "（启源）（业主）政府采购协议.docx": "（启源）（业主）政府采购协议.docx",
    "（业主）抽取专家委托书.docx": "（业主）抽取专家委托书.docx",
    "（业主）负责人授权委托书.docx": "（业主）负责人授权委托书.docx",
    "（业主）业主评标授权委托书.docx": "（业主）业主评标授权委托书.docx",
    "（监督）监督委托书.docx": "（监督）监督委托书.docx",
    "（业主）八严禁.docx": "（业主）八严禁.docx",
    "（启源）招标代理机构招标项目部人员组成情况表.xlsx": "（启源）招标代理机构招标项目部人员组成情况表.xlsx"
}

# 定义一个函数，根据给定的数据填充模板
def fill_template(template_filename, output_filename, data):
    # 根据模板文件的扩展名确定文件类型
    _, ext = os.path.splitext(template_filename)

    # 使用相应的库打开模板文件
    if ext == ".docx":
        document = Document(os.path.join(TEMPLATES_DIR, template_filename))
    elif ext == ".xlsx":
        workbook = openpyxl.load_workbook(os.path.join(TEMPLATES_DIR, template_filename))
        worksheet = workbook.active

    # 遍历数据字典，将占位符替换为对应的值
    for key, value in data.items():
        if ext == ".docx":
            # 处理Word文档中的占位符
            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if key in run.text:
                        # 保存原始格式
                        original_format = run.font._element
                        # 替换文本并设置格式
                        new_text = run.text.replace(key, value)
                        run.text = new_text
                        run.font._element = original_format
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if key in run.text:
                                    # 保存原始格式
                                    original_format = run.font._element
                                    # 替换文本并设置格式
                                    new_text = run.text.replace(key, value)
                                    run.text = new_text
                                    run.font._element = original_format
        elif ext == ".xlsx":
            # 处理Excel表格中的占位符
            for row in worksheet.iter_rows():
                for cell in row:
                    if key in str(cell.value):
                        # 替换文本并设置格式
                        new_text = str(cell.value).replace(key, value)
                        cell.value = new_text

    # 保存生成的文件
    if ext == ".docx":
        document.save(os.path.join(OUTPUT_DIR, output_filename))
    elif ext == ".xlsx":
        workbook.save(os.path.join(OUTPUT_DIR, output_filename))

# 定义填充模板文件所需的数据
data = {
    "[项目名称]": "一个神秘的项目",
    "[项目金额]": "2块5毛",
    "[项目编号]": "潜龙1号",
    "[采购方式]": "公开招标",
    "[采购类型]": "服务",
    "[资金来源]": "自筹资金",
    "[启源项目负责人]": "我",
    "[启源联系电话]": "0371-22709985",
    "[业主单位名称]": "保护伞公司",
    "[业主方项目负责人]": "奥斯威尔·E·斯宾塞",
    "[业主联系电话]": "0000-0000",
    "[项目概况]": "从事世界范围内的消灭生化武器和生化危胁的工作。",
    "[监督部门]": "浣熊市政府"
}

# 遍历模板文件，生成对应的输出文件
for template_filename, output_filename in TEMPLATE_FILENAMES.items():
    try:
        fill_template(template_filename, output_filename, data)
        print(f"{template_filename}已成功生成{output_filename}")
    except Exception as e:
        print(f"{template_filename}生成{output_filename}失败: {e}")
